import tkinter as tk
from tkinter import messagebox, filedialog
from docx import Document
from docx2pdf import convert
import os
import re


def replace_any_fields(doc, data_map):
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        new_text = original_text

        fields = re.findall(r'\[([^\]]+)\]', original_text)

        for field in fields:
            for key, value in data_map.items():
                clean_field = ' '.join(field.split())
                clean_key = ' '.join(key.replace('[', '').replace(']', '').split())

                if clean_field == clean_key and value:
                    new_text = new_text.replace(f'[{field}]', value)
                    break

        if new_text != original_text:
            paragraph.clear()
            paragraph.add_run(new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text

                    fields = re.findall(r'\[([^\]]+)\]', original_text)

                    for field in fields:
                        for key, value in data_map.items():
                            clean_field = ' '.join(field.split())
                            clean_key = ' '.join(key.replace('[', '').replace(']', '').split())

                            if clean_field == clean_key and value:
                                new_text = new_text.replace(f'[{field}]', value)
                                break

                    if new_text != original_text:
                        paragraph.clear()
                        paragraph.add_run(new_text)


def load_from_txt():
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if file_path:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                loaded_data = {}
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('#'):
                        continue
                    if '=' in line:
                        key, value = line.split('=', 1)
                        loaded_data[key.strip()] = value.strip()

            entry_city.delete(0, tk.END)
            entry_city.insert(0, loaded_data.get("City", ""))
            entry_date.delete(0, tk.END)
            entry_date.insert(0, loaded_data.get("ContractDate", ""))
            entry_owner.delete(0, tk.END)
            entry_owner.insert(0, loaded_data.get("LandlordName", ""))
            entry_position.delete(0, tk.END)
            entry_position.insert(0, loaded_data.get("Position", ""))
            entry_basis.delete(0, tk.END)
            entry_basis.insert(0, loaded_data.get("Basis", ""))
            entry_tenant.delete(0, tk.END)
            entry_tenant.insert(0, loaded_data.get("TenantName", ""))
            entry_series.delete(0, tk.END)
            entry_series.insert(0, loaded_data.get("PassportSeries", ""))
            entry_num.delete(0, tk.END)
            entry_num.insert(0, loaded_data.get("PassportNumber", ""))
            entry_passport_issued.delete(0, tk.END)
            entry_passport_issued.insert(0, loaded_data.get("PassportIssued", ""))
            entry_code.delete(0, tk.END)
            entry_code.insert(0, loaded_data.get("DepartmentCode", ""))
            entry_address.delete(0, tk.END)
            entry_address.insert(0, loaded_data.get("PropertyAddress", ""))
            entry_year_built.delete(0, tk.END)
            entry_year_built.insert(0, loaded_data.get("YearBuilt", ""))
            entry_rooms.delete(0, tk.END)
            entry_rooms.insert(0, loaded_data.get("RoomsCount", ""))
            entry_area.delete(0, tk.END)
            entry_area.insert(0, loaded_data.get("TotalArea", ""))
            entry_living_area.delete(0, tk.END)
            entry_living_area.insert(0, loaded_data.get("LivingArea", ""))
            entry_room.delete(0, tk.END)
            entry_room.insert(0, loaded_data.get("PropertyType", ""))
            entry_keys.delete(0, tk.END)
            entry_keys.insert(0, loaded_data.get("KeysCount", ""))
            entry_price.delete(0, tk.END)
            entry_price.insert(0, loaded_data.get("MonthlyPayment", ""))

            messagebox.showinfo("Успех", "Данные загружены из TXT!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить TXT: {e}")


def generate():
    data_map = {
        "вписать нужное": entry_city.get(),
        "число, месяц, год": entry_date.get(),
        "Указать наименование собственника жилого помещения или управомоченного лица": entry_owner.get(),
        "должность, Ф. И. О. полностью": entry_position.get(),
        "устава, положения, доверенности": entry_basis.get(),
        "Ф. И. О. полностью": entry_tenant.get(),
        "вписать нужное": entry_series.get(),
        "значение": entry_num.get(),
        "наименование органа, выдавшего паспорт, дата выдачи": entry_passport_issued.get(),
        "вписать нужное": entry_code.get(),
        "вписать нужное": entry_address.get(),
        "указать год постройки": entry_year_built.get(),
        "значение": entry_rooms.get(),
        "цифрами и прописью": entry_area.get(),
        "цифрами и прописью": entry_living_area.get(),
        "квартира/жилой дом/часть квартиры/часть жилого дома": entry_room.get(),
        "значение": entry_keys.get(),
        "цифрами и прописью": entry_price.get(),
    }

    required_fields = [entry_city, entry_date, entry_owner, entry_tenant, entry_address, entry_room, entry_price]
    for field in required_fields:
        if not field.get().strip():
            messagebox.showwarning("Ошибка", "Заполните все обязательные поля!")
            return

    if not os.path.exists("template.docx"):
        messagebox.showerror("Ошибка", "Файл template.docx не найден!\nПоместите его в папку со скриптом.")
        return

    try:
        doc = Document("template.docx")

        print("=== ДЕБАГ: Поиск полей в документе ===")
        for i, paragraph in enumerate(doc.paragraphs[:10]):
            text = paragraph.text
            if '[' in text and ']' in text:
                print(f"Параграф {i}: {text[:100]}...")

        replace_any_fields(doc, data_map)

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")],
            title="Сохранить договор как",
            initialfile="Договор_аренды.docx"
        )

        if save_path:
            doc.save(save_path)
            try:
                pdf_path = save_path.replace(".docx", ".pdf")
                convert(save_path, pdf_path)
                messagebox.showinfo("Готово", f"Договор успешно создан!\n\nWord: {save_path}\nPDF: {pdf_path}")
                os.startfile(save_path)
            except Exception as pdf_error:
                messagebox.showwarning("Частичный успех",
                                       f"Word файл создан: {save_path}\nНо PDF не сконвертирован: {pdf_error}")
                os.startfile(save_path)

    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при создании документа: {e}")
        print(f"Детали ошибки: {e}")


root = tk.Tk()
root.title("Генератор договоров аренды жилья")
root.geometry("650x900")


def create_field(parent, label_text):
    frame = tk.Frame(parent)
    frame.pack(fill="x", pady=5)

    label = tk.Label(frame, text=label_text, width=25, anchor="w")
    label.pack(side="left", padx=(0, 10))

    entry = tk.Entry(frame, width=40)
    entry.pack(side="left")

    return entry


entry_city = create_field(root, "Город:")
entry_date = create_field(root, "Дата (чч.мм.гггг):")
entry_owner = create_field(root, "ФИО Наймодателя:")
entry_position = create_field(root, "Должность Наймодателя:")
entry_basis = create_field(root, "Основание действия:")
entry_tenant = create_field(root, "ФИО Нанимателя:")
entry_series = create_field(root, "Серия паспорта:")
entry_num = create_field(root, "Номер паспорта:")
entry_passport_issued = create_field(root, "Паспорт выдан:")
entry_code = create_field(root, "Код подразделения:")
entry_address = create_field(root, "Адрес помещения:")
entry_year_built = create_field(root, "Год постройки:")
entry_rooms = create_field(root, "Количество комнат:")
entry_area = create_field(root, "Общая площадь:")
entry_living_area = create_field(root, "Жилая площадь:")
entry_room = create_field(root, "Тип помещения:")
entry_keys = create_field(root, "Количество ключей:")
entry_price = create_field(root, "Ежемесячная плата:")

btn_frame = tk.Frame(root)
btn_frame.pack(pady=20)

btn_load = tk.Button(btn_frame, text="Загрузить из TXT", command=load_from_txt,
                     bg="#2196F3", fg="white", font=("Arial", 10), width=15)
btn_load.pack(side="left", padx=5)

btn_generate = tk.Button(btn_frame, text="Создать договор", command=generate,
                         bg="#4CAF50", fg="white", font=("Arial", 10, "bold"), width=15)
btn_generate.pack(side="left", padx=5)

info_label = tk.Label(root,
                      text="Убедитесь, что template.docx в папке со скриптом",
                      fg="blue", font=("Arial", 9))
info_label.pack(pady=10)


def debug_template():
    try:
        doc = Document("template.docx")
        print("=== СОДЕРЖИМОЕ ШАБЛОНА ===")
        for i, paragraph in enumerate(doc.paragraphs[:20]):
            text = paragraph.text.strip()
            if text:
                print(f"{i:3}: {text}")
                for run in paragraph.runs:
                    if '[' in run.text and ']' in run.text:
                        print(f"     Run: '{run.text}'")
    except Exception as e:
        print(f"Ошибка при чтении шаблона: {e}")


btn_debug = tk.Button(root, text="Отладка шаблона", command=debug_template,
                      bg="#FF9800", fg="white", font=("Arial", 8))
btn_debug.pack(pady=5)

root.mainloop()